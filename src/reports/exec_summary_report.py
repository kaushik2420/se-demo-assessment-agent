"""
Generate the CEO-facing monthly executive summary (.docx).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


SS_BLUE = RGBColor(0x29, 0x6F, 0xC4)
SS_DARK = RGBColor(0x1A, 0x2B, 0x3D)
SS_GREEN = RGBColor(0x2E, 0xA0, 0x5F)
SS_RED = RGBColor(0xC9, 0x35, 0x35)
SS_GREY = RGBColor(0x6B, 0x72, 0x80)


def _set_cell(cell, text, bold=False, color=None, size=10):
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def _h(doc, text, lvl=1, color=SS_DARK):
    h = doc.add_heading(text, level=lvl)
    for r in h.runs:
        r.font.color.rgb = color


def _p(doc, text, bold=False, size=11, color=None, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color


def build_exec_summary(
    month_label: str, summary: dict, output_path: str
) -> str:
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.7); s.bottom_margin = Inches(0.7)
        s.left_margin = Inches(0.8); s.right_margin = Inches(0.8)

    # Title
    title = doc.add_paragraph()
    tr = title.add_run("CEO Executive Summary — Solution Engineering")
    tr.bold = True; tr.font.size = Pt(20); tr.font.color.rgb = SS_DARK
    sub = doc.add_paragraph()
    sr = sub.add_run(f"{month_label} · for the CEO  ·  prepared by kaushik")
    sr.font.size = Pt(11); sr.font.color.rgb = SS_GREY

    # Headline
    _p(doc, summary["headline"], bold=True, size=13, color=SS_BLUE, italic=True)

    # Metric strip
    m = summary["month_metrics"]
    tbl = doc.add_table(rows=2, cols=5)
    headers = ["Avg SE Score", "Best SE", "Worst SE", "Feature-Selling %", "AE Interrupts / Call"]
    values = [
        f"{m['avg_se_score']}/5",
        f"{m['best_se']['name']} ({m['best_se']['score']})",
        f"{m['worst_se']['name']} ({m['worst_se']['score']})",
        f"{int(m['feature_selling_pct_of_demos']*100)}%",
        f"{m['ae_interruption_avg_per_call']}",
    ]
    for i, h in enumerate(headers):
        _set_cell(tbl.rows[0].cells[i], h, bold=True, size=9, color=SS_GREY)
        _set_cell(tbl.rows[1].cells[i], values[i], bold=True, size=12, color=SS_DARK)

    doc.add_paragraph()

    # CX maturity mix
    _h(doc, "CX maturity mix of prospects this month", lvl=2)
    cx = m["cx_maturity_mix"]
    cx_tbl = doc.add_table(rows=1, cols=len(cx) or 1)
    cx_tbl.style = "Light List Accent 1"
    for i, (cat, frac) in enumerate(cx.items()):
        _set_cell(cx_tbl.rows[0].cells[i], f"{cat}\n{int(frac*100)}%", bold=True, size=10)

    # PRODUCT GAPS
    _h(doc, "Top 5 product gaps (engineering should care)", lvl=2)
    pg = doc.add_table(rows=1, cols=4)
    pg.style = "Light Grid Accent 1"
    for i, h in enumerate(["#", "Gap", "Evidence calls", "$ at risk"]):
        _set_cell(pg.rows[0].cells[i], h, bold=True, size=10)
    for g in summary["top_5_product_gaps"]:
        row = pg.add_row().cells
        _set_cell(row[0], g["rank"], bold=True, size=10, color=SS_BLUE)
        _set_cell(row[1], g["gap"], size=10)
        _set_cell(row[2], ", ".join(g["evidence_calls"][:2]) if g["evidence_calls"] else "—", size=9, color=SS_GREY)
        _set_cell(row[3], g["deals_at_risk_estimate"], size=10, color=SS_RED)

    # PROCESS GAPS
    _h(doc, "Top 5 process gaps (GTM should own)", lvl=2)
    pr = doc.add_table(rows=1, cols=4)
    pr.style = "Light Grid Accent 1"
    for i, h in enumerate(["#", "Gap", "Owner", "Evidence calls"]):
        _set_cell(pr.rows[0].cells[i], h, bold=True, size=10)
    for g in summary["top_5_process_gaps"]:
        row = pr.add_row().cells
        _set_cell(row[0], g["rank"], bold=True, size=10, color=SS_BLUE)
        _set_cell(row[1], g["gap"], size=10)
        _set_cell(row[2], g["recommended_owner"], size=10, color=SS_DARK, bold=True)
        _set_cell(row[3], ", ".join(g["evidence_calls"][:2]) if g["evidence_calls"] else "—", size=9, color=SS_GREY)

    # AE QUALITY
    _h(doc, "Account Executive quality risks", lvl=2)
    if not summary.get("ae_quality_risks"):
        _p(doc, "No AE quality patterns flagged this month.", size=11, color=SS_GREY)
    for ae in summary["ae_quality_risks"]:
        _p(doc, f"{ae['ae_name']}", bold=True, size=12, color=SS_RED)
        _p(doc, f"Pattern:  {ae['pattern']}", size=11)
        _p(doc, f"Recommendation:  {ae['recommendation']}", size=11, italic=True, color=SS_BLUE)
        doc.add_paragraph()

    # COMPETITIVE
    _h(doc, "Competitive intelligence", lvl=2)
    _p(doc, summary["competitive_intel"]["where_we_lose"], size=11, italic=True)
    if summary["competitive_intel"]["most_mentioned_competitors"]:
        ct = doc.add_table(rows=1, cols=3)
        ct.style = "Light Grid Accent 1"
        for i, h in enumerate(["Competitor", "Mentions this month", "Approx. win-rate proxy"]):
            _set_cell(ct.rows[0].cells[i], h, bold=True, size=10)
        for c in summary["competitive_intel"]["most_mentioned_competitors"]:
            row = ct.add_row().cells
            _set_cell(row[0], c["name"], size=10, bold=True)
            _set_cell(row[1], c["mentions"], size=10)
            _set_cell(row[2], f"{int(c['win_rate_proxy']*100)}%", size=10)

    # CEO ACTIONS
    _h(doc, "Top 5 actions for the CEO", lvl=2, color=SS_BLUE)
    for i, action in enumerate(summary["ceo_top_5_actions"], start=1):
        p = doc.add_paragraph()
        n = p.add_run(f"{i}.  "); n.bold = True; n.font.color.rgb = SS_BLUE; n.font.size = Pt(12)
        t = p.add_run(action); t.font.size = Pt(11)

    # Footer
    doc.add_paragraph()
    foot = doc.add_paragraph()
    fr = foot.add_run(
        "Generated by SurveySparrow SE Demo Assessment Agent · "
        f"Based on {summary.get('month_metrics', {}).get('avg_se_score', '?')}-class demos analyzed via Recall.ai + Claude."
    )
    fr.font.size = Pt(9); fr.italic = True; fr.font.color.rgb = SS_GREY

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path
