"""
Generate the per-SE monthly Word report (.docx) using python-docx.

Sections:
  1. Header card: name, month, final-score, percentile vs SaaS industry
  2. Score table: criterion × score × industry median × gap
  3. Trend line (vs last 3 months)
  4. Top 3 strengths
  5. Top 3 gaps
  6. Coaching action of the month
  7. Notable call moments (best & worst evidence quotes)
  8. CX maturity & deal-intelligence rollup
"""

from __future__ import annotations

from pathlib import Path
from statistics import mean
from typing import List

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.utils.benchmarks import INDUSTRY_MEDIAN_BY_CRITERION, percentile_of
from src.utils.rubric import RUBRIC


SS_BLUE = RGBColor(0x29, 0x6F, 0xC4)
SS_DARK = RGBColor(0x1A, 0x2B, 0x3D)
SS_GREEN = RGBColor(0x2E, 0xA0, 0x5F)
SS_RED = RGBColor(0xC9, 0x35, 0x35)
SS_GREY = RGBColor(0x6B, 0x72, 0x80)


def _set_cell(cell, text, bold=False, color=None, size=10, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def _add_heading(doc, text, level=1, color=SS_DARK):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color


def _add_para(doc, text, bold=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def _score_color(score, median):
    if score >= median + 0.4:
        return SS_GREEN
    if score <= median - 0.4:
        return SS_RED
    return SS_DARK


def build_se_report(
    se_name: str,
    se_email: str,
    month_label: str,
    scorecards: List[dict],
    insights: List[dict],
    last_3_months_finals: List[float],
    output_path: str,
) -> str:
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # === HEADER ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Solution Engineer Monthly Coaching Report")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = SS_DARK

    sub = doc.add_paragraph()
    sub_run = sub.add_run(f"{se_name}  ·  {month_label}")
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = SS_GREY

    avg_final = round(mean(s["weighted_final"] for s in scorecards), 2) if scorecards else 0
    percentile = percentile_of(avg_final)

    # === HEADLINE CARD ===
    card = doc.add_table(rows=1, cols=3)
    card.autofit = False
    for i, w in enumerate([2.5, 2.5, 2.5]):
        card.columns[i].width = Inches(w)
    cells = card.rows[0].cells
    _set_cell(cells[0], "FINAL SCORE", bold=True, size=9, color=SS_GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(cells[1], "CALLS ANALYZED", bold=True, size=9, color=SS_GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(cells[2], "INDUSTRY PERCENTILE", bold=True, size=9, color=SS_GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
    card2 = doc.add_table(rows=1, cols=3)
    for i, w in enumerate([2.5, 2.5, 2.5]):
        card2.columns[i].width = Inches(w)
    c2 = card2.rows[0].cells
    _set_cell(c2[0], f"{avg_final} / 5.0", bold=True, size=22, color=SS_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(c2[1], str(len(scorecards)), bold=True, size=22, color=SS_DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(c2[2], f"P{percentile}", bold=True, size=22,
              color=SS_GREEN if percentile >= 75 else (SS_RED if percentile <= 25 else SS_DARK),
              align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # === HOW YOU STACK UP VS SAAS INDUSTRY ===
    _add_heading(doc, "How you stack up vs the SaaS SE industry", level=2)
    _add_para(doc,
        f"Across {len(scorecards)} demos this month, your weighted score is {avg_final}/5 — "
        f"approximately the {percentile}th percentile of B2B SaaS Solution Engineers "
        f"(industry medians from Gartner SE Excellence, PreSales Collective, and Bain SaaS GTM benchmarks).",
        size=11)

    # === CRITERION TABLE ===
    _add_heading(doc, "Score by criterion", level=2)
    per_crit_avg = {}
    for c in RUBRIC:
        vals = [sc["per_criterion_score"].get(c.name, 0) for sc in scorecards]
        per_crit_avg[c.name] = round(mean(vals), 2) if vals else 0

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Criterion", "Weight", "Your score", "Industry median", "Gap"]):
        _set_cell(hdr[i], h, bold=True, size=10, color=SS_DARK)
    for c in RUBRIC:
        row = tbl.add_row().cells
        median = INDUSTRY_MEDIAN_BY_CRITERION[c.name]
        score = per_crit_avg[c.name]
        gap = round(score - median, 2)
        _set_cell(row[0], c.name, size=10)
        _set_cell(row[1], f"{c.weight:.0f}%", size=10)
        _set_cell(row[2], f"{score:.2f}", size=10, bold=True, color=_score_color(score, median))
        _set_cell(row[3], f"{median:.2f}", size=10, color=SS_GREY)
        _set_cell(row[4], f"{'+' if gap >= 0 else ''}{gap:.2f}", size=10,
                  color=SS_GREEN if gap >= 0 else SS_RED, bold=True)

    doc.add_paragraph()

    # === TREND ===
    _add_heading(doc, "Your last 3 months", level=2)
    trend_text = " → ".join(f"{x:.2f}" for x in last_3_months_finals) if last_3_months_finals else "First month tracked"
    arrow = "↑ improving" if len(last_3_months_finals) >= 2 and last_3_months_finals[-1] > last_3_months_finals[0] \
            else ("↓ declining" if len(last_3_months_finals) >= 2 and last_3_months_finals[-1] < last_3_months_finals[0] else "stable")
    _add_para(doc, f"{trend_text}    ({arrow})", size=11, bold=True,
              color=SS_GREEN if "improving" in arrow else (SS_RED if "declining" in arrow else SS_DARK))

    # === STRENGTHS / GAPS / COACHING ===
    all_strengths, all_gaps, all_actions = [], [], []
    for sc in scorecards:
        q = sc.get("qualitative", {})
        all_strengths.extend(q.get("top_3_strengths", []))
        all_gaps.extend(q.get("top_3_gaps", []))
        if q.get("one_coaching_action"):
            all_actions.append(q["one_coaching_action"])

    _add_heading(doc, "Your top 3 strengths this month", level=2)
    for s in all_strengths[:3]:
        doc.add_paragraph(s, style="List Bullet")

    _add_heading(doc, "Top 3 areas of improvement", level=2)
    for g in all_gaps[:3]:
        doc.add_paragraph(g, style="List Bullet")

    _add_heading(doc, "Your single coaching action for next month", level=2)
    if all_actions:
        p = doc.add_paragraph()
        r = p.add_run(all_actions[0])
        r.italic = True
        r.font.size = Pt(11)
        r.font.color.rgb = SS_BLUE

    # === DEAL-INTEL ROLLUP ===
    _add_heading(doc, "What your demos told us about the market", level=2)
    cx_mix = {}
    for i in insights:
        cat = i["cx_maturity"]["category"]
        cx_mix[cat] = cx_mix.get(cat, 0) + 1
    feature_sellers = sum(1 for i in insights if i["se_selling_style"]["verdict"] == "feature_seller")
    interrupts = [i["ae_behavior"]["interruption_count"] for i in insights]

    intel = doc.add_table(rows=4, cols=2)
    intel.style = "Light List Accent 1"
    rows = intel.rows
    _set_cell(rows[0].cells[0], "CX maturity mix", bold=True, size=10)
    _set_cell(rows[0].cells[1], ", ".join(f"{k}: {v}" for k, v in cx_mix.items()) or "—", size=10)
    _set_cell(rows[1].cells[0], "Demos classified feature-selling", bold=True, size=10)
    _set_cell(rows[1].cells[1], f"{feature_sellers} of {len(insights)}", size=10,
              color=SS_RED if feature_sellers > len(insights)/2 else SS_DARK)
    _set_cell(rows[2].cells[0], "Avg AE interruptions per call", bold=True, size=10)
    _set_cell(rows[2].cells[1], f"{round(mean(interrupts),1) if interrupts else 0}", size=10)
    all_competitors = []
    for i in insights:
        all_competitors.extend(c["name"] for c in i["competitors_mentioned"])
    _set_cell(rows[3].cells[0], "Competitors you faced", bold=True, size=10)
    _set_cell(rows[3].cells[1], ", ".join(sorted(set(all_competitors))) or "—", size=10)

    # === EVIDENCE FOOTER ===
    doc.add_paragraph()
    _add_heading(doc, "Evidence: highest- and lowest-scoring moments", level=2)
    # find best & worst sub from first call's scores
    if scorecards:
        flat = []
        for sc in scorecards:
            for crit, subs in sc["scores"].items():
                for sub, payload in subs.items():
                    flat.append((payload["score"], crit, sub, payload.get("evidence", "")))
        flat.sort(reverse=True)
        if flat:
            best = flat[0]; worst = flat[-1]
            _add_para(doc, f"BEST  ·  {best[1]} → {best[2]}  ({best[0]}/5)", bold=True, size=10, color=SS_GREEN)
            _add_para(doc, f"  \"{best[3]}\"", size=10, color=SS_GREY)
            _add_para(doc, f"WORST  ·  {worst[1]} → {worst[2]}  ({worst[0]}/5)", bold=True, size=10, color=SS_RED)
            _add_para(doc, f"  \"{worst[3]}\"", size=10, color=SS_GREY)

    # Footer note
    doc.add_paragraph()
    foot = doc.add_paragraph()
    fr = foot.add_run(
        "Generated by SurveySparrow SE Demo Assessment Agent · "
        "Questions? Reply to this email (kaushikn2416@gmail.com is on CC)."
    )
    fr.font.size = Pt(9)
    fr.font.color.rgb = SS_GREY
    fr.italic = True

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path
