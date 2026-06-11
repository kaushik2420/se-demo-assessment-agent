"""
Changelog router — admin + manager only.

Endpoints:
  GET    /changelog              list all entries (newest entry_number first)
  POST   /changelog              create a new entry
  GET    /changelog/{id}         single entry detail
  PATCH  /changelog/{id}         update entry
  DELETE /changelog/{id}         delete entry
  GET    /changelog/export.md    download all entries as markdown

The entry_number auto-increments based on the current max — admins don't
have to assign numbers manually unless they want to (a number can still be
passed explicitly).
"""

from __future__ import annotations

from datetime import datetime, timezone
from typing import List, Optional

from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy import desc, func
from sqlalchemy.orm import Session

from app.db import get_db
from app.deps import CurrentUser, get_current_user, require_role
from app.models import ChangelogEntry


router = APIRouter()


# ─────────────────────────────────────────────────────────────────────────
# Schemas
# ─────────────────────────────────────────────────────────────────────────

class ChangelogItem(BaseModel):
    id: int
    entry_number: int
    title: str
    issue: str
    rca: str
    fix: str
    status: str
    entry_date: str
    created_by: Optional[str]
    created_at: str
    updated_at: Optional[str]
    updated_by: Optional[str]


class ChangelogCreate(BaseModel):
    title: str
    issue: str
    rca: str
    fix: str
    status: str = "shipped"
    entry_date: Optional[str] = None      # ISO "YYYY-MM-DD"; defaults to today
    entry_number: Optional[int] = None    # defaults to max+1


class ChangelogPatch(BaseModel):
    title: Optional[str] = None
    issue: Optional[str] = None
    rca: Optional[str] = None
    fix: Optional[str] = None
    status: Optional[str] = None
    entry_date: Optional[str] = None
    entry_number: Optional[int] = None


# ─────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────

VALID_STATUSES = ("shipped", "pending", "blocked", "deferred")


def _parse_date(s: Optional[str]) -> Optional[datetime]:
    if not s:
        return None
    try:
        from datetime import date as _d
        d = _d.fromisoformat(str(s)[:10])
        return datetime(d.year, d.month, d.day, tzinfo=timezone.utc)
    except Exception:
        return None


def _to_item(row: ChangelogEntry) -> ChangelogItem:
    return ChangelogItem(
        id=row.id,
        entry_number=row.entry_number,
        title=row.title,
        issue=row.issue,
        rca=row.rca,
        fix=row.fix,
        status=row.status or "shipped",
        entry_date=row.entry_date.isoformat() if row.entry_date else "",
        created_by=row.created_by,
        created_at=row.created_at.isoformat() if row.created_at else "",
        updated_at=row.updated_at.isoformat() if row.updated_at else None,
        updated_by=row.updated_by,
    )


# ─────────────────────────────────────────────────────────────────────────
# Endpoints
# ─────────────────────────────────────────────────────────────────────────

@router.get("", response_model=List[ChangelogItem],
            dependencies=[Depends(require_role("admin", "manager"))])
def list_entries(db: Session = Depends(get_db)):
    """All entries, newest entry_number first."""
    rows = db.query(ChangelogEntry).order_by(desc(ChangelogEntry.entry_number)).all()
    return [_to_item(r) for r in rows]


@router.post("/reseed", dependencies=[Depends(require_role("admin"))])
def reseed(wipe: bool = False):
    """Admin rescue endpoint — re-import entries from the bundled markdown
    snapshot. By default tops up only missing entries (won't overwrite anything
    already in the DB). Pass `?wipe=true` to delete every row first, then
    re-import — use this to fully reset to the snapshot state."""
    from app.services.changelog_seeder import force_reseed
    return force_reseed(wipe_existing=wipe)


@router.get("/export.docx", dependencies=[Depends(require_role("admin", "manager"))])
def export_docx(db: Session = Depends(get_db)):
    """Download the entire changelog as a .docx file with proper formatting.
    Rishul / leadership can attach this to email or paste into board decks."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from io import BytesIO

    rows = db.query(ChangelogEntry).order_by(desc(ChangelogEntry.entry_number)).all()

    doc = Document()

    # Title
    t = doc.add_heading("SE Coach — Post-Deployment Change Log", level=0)
    for run in t.runs:
        run.font.color.rgb = RGBColor(0x25, 0x30, 0x43)  # ss-navy

    p = doc.add_paragraph()
    p.add_run("A running record of every bug, feedback item, and feature request raised by the team after the day-1 deployment.").italic = True

    # Meta
    meta = doc.add_paragraph()
    meta_run = meta.add_run(
        f"Exported: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}  ·  {len(rows)} entries"
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x5A, 0x6B, 0x85)  # ss-navy-soft

    doc.add_paragraph()  # spacer

    for r in rows:
        # Entry heading: "#34 — Title"
        h = doc.add_heading(f"#{r.entry_number} — {r.title}", level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x25, 0x30, 0x43)

        # Status + Date line
        meta_para = doc.add_paragraph()
        status_run = meta_para.add_run(f"Status: {r.status or 'shipped'}")
        status_run.bold = True
        status_run.font.size = Pt(10)
        date_str = r.entry_date.strftime("%Y-%m-%d") if r.entry_date else "—"
        date_run = meta_para.add_run(f"     ·     Date: {date_str}")
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(0x5A, 0x6B, 0x85)

        # Issue
        h = doc.add_paragraph()
        run = h.add_run("Issue / Feedback:")
        run.bold = True
        run.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)  # amber
        doc.add_paragraph(r.issue)

        # RCA
        h = doc.add_paragraph()
        run = h.add_run("RCA:")
        run.bold = True
        run.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)
        doc.add_paragraph(r.rca)

        # Fix
        h = doc.add_paragraph()
        run = h.add_run("Fix:")
        run.bold = True
        run.font.color.rgb = RGBColor(0x1B, 0x7E, 0x47)  # emerald
        doc.add_paragraph(r.fix)

        # Separator
        sep = doc.add_paragraph()
        sep_run = sep.add_run("—" * 60)
        sep_run.font.color.rgb = RGBColor(0xC8, 0xE8, 0xF2)  # cyan-soft
        sep_run.font.size = Pt(8)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition":
                f'attachment; filename="se-coach-changelog-{datetime.utcnow():%Y%m%d}.docx"',
        },
    )


@router.get("/export.md", dependencies=[Depends(require_role("admin", "manager"))])
def export_markdown(db: Session = Depends(get_db)):
    """Download the entire changelog as a markdown file. Same shape as the
    original POST_DEPLOYMENT_CHANGELOG.md so it can be pasted into Google
    Docs or attached to a Slack thread."""
    rows = db.query(ChangelogEntry).order_by(desc(ChangelogEntry.entry_number)).all()

    parts = [
        "# SE Coach — Post-Deployment Change Log",
        "",
        "A running record of every bug, feedback item, and feature request raised by the team after the day-1 deployment.",
        "",
        f"Exported: **{datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}** · "
        f"{len(rows)} entries",
        "",
        "---",
        "",
    ]
    for r in rows:
        date_str = r.entry_date.strftime("%Y-%m-%d") if r.entry_date else "—"
        parts.append(f"## #{r.entry_number} — {r.title}")
        parts.append("")
        parts.append(f"**Issue / Feedback:** {r.issue}")
        parts.append("")
        parts.append(f"**RCA:** {r.rca}")
        parts.append("")
        parts.append(f"**Fix:** {r.fix}")
        parts.append("")
        parts.append(f"**Date:** {date_str}")
        if r.status and r.status != "shipped":
            parts.append("")
            parts.append(f"**Status:** {r.status}")
        parts.append("")
        parts.append("---")
        parts.append("")

    md = "\n".join(parts)
    return StreamingResponse(
        iter([md]),
        media_type="text/markdown",
        headers={
            "Content-Disposition":
                f'attachment; filename="se-coach-changelog-{datetime.utcnow():%Y%m%d}.md"',
        },
    )


@router.get("/{entry_id}", response_model=ChangelogItem,
            dependencies=[Depends(require_role("admin", "manager"))])
def get_entry(entry_id: int, db: Session = Depends(get_db)):
    row = db.query(ChangelogEntry).filter(ChangelogEntry.id == entry_id).first()
    if not row:
        raise HTTPException(404, "Entry not found")
    return _to_item(row)


@router.post("", response_model=ChangelogItem,
             status_code=status.HTTP_201_CREATED,
             dependencies=[Depends(require_role("admin", "manager"))])
def create_entry(
    req: ChangelogCreate,
    actor: CurrentUser = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if req.status and req.status not in VALID_STATUSES:
        raise HTTPException(400, f"Invalid status {req.status!r}; must be one of {VALID_STATUSES}")

    # Auto-assign next entry number if not provided
    next_num = req.entry_number
    if next_num is None:
        current_max = db.query(func.max(ChangelogEntry.entry_number)).scalar() or 0
        next_num = current_max + 1
    else:
        # Duplicate-number guard
        if db.query(ChangelogEntry).filter(ChangelogEntry.entry_number == next_num).first():
            raise HTTPException(409, f"Entry #{next_num} already exists")

    entry_date = _parse_date(req.entry_date) or datetime.now(timezone.utc)

    row = ChangelogEntry(
        entry_number=next_num,
        title=req.title.strip(),
        issue=req.issue.strip(),
        rca=req.rca.strip(),
        fix=req.fix.strip(),
        status=req.status,
        entry_date=entry_date,
        created_by=actor.email,
    )
    db.add(row)
    db.commit()
    db.refresh(row)
    print(f"[changelog] created entry #{row.entry_number} by {actor.email!r}: {row.title!r}")
    return _to_item(row)


@router.patch("/{entry_id}", response_model=ChangelogItem,
              dependencies=[Depends(require_role("admin", "manager"))])
def patch_entry(
    entry_id: int,
    patch: ChangelogPatch,
    actor: CurrentUser = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    row = db.query(ChangelogEntry).filter(ChangelogEntry.id == entry_id).first()
    if not row:
        raise HTTPException(404, "Entry not found")

    data = patch.model_dump(exclude_unset=True)
    if "status" in data and data["status"] not in VALID_STATUSES:
        raise HTTPException(400, f"Invalid status {data['status']!r}")
    if "entry_number" in data and data["entry_number"] is not None:
        # Duplicate-number guard if changing the number
        dupe = (db.query(ChangelogEntry)
                .filter(ChangelogEntry.entry_number == data["entry_number"])
                .filter(ChangelogEntry.id != row.id).first())
        if dupe:
            raise HTTPException(409, f"Entry #{data['entry_number']} already exists")

    for k in ("title", "issue", "rca", "fix", "status", "entry_number"):
        if k in data and data[k] is not None:
            v = data[k]
            if isinstance(v, str):
                v = v.strip()
            setattr(row, k, v)
    if "entry_date" in data:
        d = _parse_date(data["entry_date"])
        if d:
            row.entry_date = d

    row.updated_at = datetime.now(timezone.utc)
    row.updated_by = actor.email
    db.commit()
    db.refresh(row)
    print(f"[changelog] updated entry #{row.entry_number} by {actor.email!r}")
    return _to_item(row)


@router.delete("/{entry_id}", status_code=204,
               dependencies=[Depends(require_role("admin"))])
def delete_entry(
    entry_id: int,
    actor: CurrentUser = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Admin only — managers can edit but not delete (audit trail discipline)."""
    row = db.query(ChangelogEntry).filter(ChangelogEntry.id == entry_id).first()
    if not row:
        raise HTTPException(404, "Entry not found")
    print(f"[changelog] DELETED entry #{row.entry_number} ({row.title!r}) by {actor.email!r}")
    db.delete(row)
    db.commit()
    return None
